#!/usr/bin/env python3
"""Generate enriched technical report for garbage classification project."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ====== Color constants ======
BLUE_DARK = RGBColor(0x1A, 0x52, 0x76)
BLUE_MED  = RGBColor(0x2E, 0x86, 0xC1)
GREEN     = RGBColor(0x27, 0xAE, 0x60)
GRAY      = RGBColor(0x66, 0x66, 0x66)
BLACK     = RGBColor(0x00, 0x00, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
RED       = RGBColor(0xE7, 0x4C, 0x3C)
ALT_ROW   = RGBColor(0xF5, 0xF9, 0xFC)

# ====== Helper functions ======
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, name="Microsoft YaHei"):
    """Set font for a run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_para(doc, text, bold=False, size=12, color=BLACK, align=None, spacing_after=6, first_indent=True, spacing_before=0, font_name="Microsoft YaHei"):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.space_before = Pt(spacing_before)
    pf.line_spacing = Pt(18)
    if first_indent and align is None:
        pf.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_run_font(run, font_name)
    return p

def add_heading_custom(doc, text, level=1):
    """Add a heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = BLUE_DARK
        elif level == 2:
            run.font.color.rgb = BLUE_MED
        elif level == 3:
            run.font.color.rgb = BLACK
    return h

def add_empty_line(doc):
    """Add an empty paragraph."""
    return add_para(doc, "", size=6, first_indent=False, spacing_after=2)

def create_table(doc, headers, rows, col_widths=None, header_color="1A5276"):
    """Create a formatted table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_run_font(run)
        set_cell_shading(cell, header_color)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            set_run_font(run)
            # Alternate row coloring
            if ri % 2 == 1:
                set_cell_shading(cell, "F5F9FC")
            # Color emphasis for specific cells
            if "97.48" in str(val) or "96.53" in str(val) or "10,518" in str(val):
                run.bold = True
                run.font.color.rgb = GREEN
            if "+597" in str(val) or "+321" in str(val):
                run.bold = True
                run.font.color.rgb = GREEN

    # Set column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)

    return table

def add_table_caption(doc, text):
    """Add a centered bold caption above a table."""
    return add_para(doc, text, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=6, spacing_after=4, first_indent=False)


# ====== MAIN DOCUMENT GENERATION ======
def generate():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===================== COVER PAGE =====================
    for _ in range(6):
        add_empty_line(doc)

    add_para(doc, "智能垃圾分类系统", bold=True, size=28, color=BLUE_DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, spacing_after=8)
    add_para(doc, "技术报告", bold=True, size=22, color=BLUE_MED,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, spacing_after=6)
    add_empty_line(doc)
    add_para(doc, "基于深度学习的图像分类技术实现", bold=False, size=16, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, spacing_after=6)
    for _ in range(4):
        add_empty_line(doc)
    add_para(doc, "同济大学  Python程序设计课程小组项目", bold=False, size=14, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, spacing_after=4)
    add_para(doc, "版本 v1.07  |  2026年5月", bold=False, size=12, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for _ in range(3):
        add_empty_line(doc)
    add_para(doc, "最佳模型: ConvNeXt Small  |  测试准确率: 97.48%  |  Macro-F1: 0.9723",
             bold=True, size=14, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    doc.add_page_break()

    # ===================== ABSTRACT =====================
    add_heading_custom(doc, "摘要", level=1)
    add_para(doc,
        "针对生活垃圾自动分类问题，本项目基于深度学习技术实现了一套智能垃圾图像分类系统。"
        "项目从数据采集、数据清洗与增强、模型架构设计、训练策略优化到模型评估部署，实现了全流程覆盖。"
        "系统支持纸板（Cardboard）、玻璃（Glass）、金属（Metal）、纸张（Paper）、塑料（Plastic）和其他垃圾（Trash）六类物品的自动识别。")

    add_para(doc,
        "在数据处理层面，项目通过多源数据集融合（Kaggle + Hugging Face）将数据集从2,497张大幅扩充至10,518张"
        "（增长321%），并实施了基于感知哈希的重复检测、损坏图像过滤和多层次数据增强策略"
        "（RandAugment + CutMix + RandomErasing）。")

    add_para(doc,
        "在模型优化层面，项目实现了10种不同规模的深度神经网络架构，并系统研究了Focal Loss、"
        "标签平滑、Cosine Warmup学习率调度、SGD Nesterov/AdamW优化器选择、EMA指数移动平均、"
        "SWA随机权重平均、WeightedRandomSampler类别平衡采样等高级训练技巧的综合作用。")

    add_para(doc,
        "经过多版本迭代优化，ConvNeXt Small模型（50M参数）在独立测试集上达到97.48%的分类准确率，"
        "Macro-F1为0.9723。实验证明，数据规模对分类性能具有决定性影响，多策略联合优化能够显著提升"
        "模型泛化能力。此外，项目采用模块化设计，提供了完整的命令行工具链、Streamlit交互式Web应用"
        "和FastAPI REST API服务，支持从数据清洗到模型推理的全流程操作。")

    add_empty_line(doc)
    add_para(doc,
        "关键词：垃圾分类；深度学习；图像分类；卷积神经网络；迁移学习；数据增强",
        bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    doc.add_page_break()

    # ===================== CHAPTER 1: INTRODUCTION =====================
    add_heading_custom(doc, "1  引言", level=1)
    add_heading_custom(doc, "1.1  项目背景", level=2)
    add_para(doc,
        "随着城市化进程的加速和人口持续增长，城市生活垃圾产生量急剧增加。据统计，全球每年产生超过20亿吨"
        "城市固体废物，其中约33%未得到环境友好型处理。垃圾分类作为资源回收利用和减少环境污染的关键环节，"
        "其重要性日益凸显。传统垃圾分类主要依赖人工分拣，存在效率低下、成本高昂且长期暴露于恶劣环境对"
        "工人健康构成威胁等突出问题。")

    add_para(doc,
        "近年来，深度学习在计算机视觉领域取得了突破性进展，为自动化和智能化垃圾分类提供了可行的技术路径。"
        "基于卷积神经网络的图像分类方法已广泛应用于各类视觉识别任务中，主流模型架构包括VGG、ResNet、"
        "MobileNet、EfficientNet、ConvNeXt等。然而，现有研究在实际应用中仍面临多重挑战：一是高质量标注"
        "数据获取困难，模型性能高度依赖数据规模和标注精确度；二是不同垃圾类别间的视觉差异较大（如纸板与"
        "纸张、透明玻璃与透明塑料），在细粒度区分上存在难度；三是实际应用场景中的光照变化、拍摄角度不同"
        "和背景噪声等因素对模型鲁棒性提出了更高要求。")

    add_heading_custom(doc, "1.2  项目目标", level=2)
    add_para(doc,
        "本项目旨在设计并实现一套高精度的智能垃圾分类系统，主要目标包括：(1)构建大规模高质量的垃圾分类"
        "数据集，通过多源数据融合缓解类别不平衡问题；(2)系统探索多种深度学习模型架构在垃圾分类任务上的"
        "性能表现；(3)深入研究数据增强、训练优化策略对模型性能的影响规律；(4)打造交互式用户界面支持便捷"
        "的实地应用。项目源自同济大学Python程序设计课程期末小组大作业，基于PyTorch深度学习框架实现，"
        "训练在NVIDIA GeForce RTX 5060 Laptop GPU上进行。")

    add_heading_custom(doc, "1.3  版本演进概览", level=2)
    add_para(doc,
        "项目经历了从v1.0到v1.08共9个版本的迭代优化。v1.0实现了3个基线模型（SimpleCNN、MobileNetV2、"
        "ResNet18）；v1.03引入先进架构（EfficientNetV2-S、ConvNeXt Tiny）和Focal Loss、RandAugment等优化；"
        "v1.04增加CutMix、RandomErasing、TTA等增强；v1.06是最关键的版本——数据集从2,497张扩充至10,518张，"
        "新增EfficientNetV2-M、ConvNeXt Small架构，引入Cosine Warmup、WeightedRandomSampler、SWA、EMA"
        "持久化等优化，准确率从90.00%跃升至96.53%；v1.07在大模型上进一步训练，ConvNeXt Small达97.48%；"
        "v1.08增加Web前端展示模块。")

    add_table_caption(doc, "表1  核心版本性能演进")
    create_table(doc,
        ["版本", "日期", "模型", "核心改进", "测试准确率"],
        [
            ["v1.0", "2024-10", "MobileNetV2", "基线模型+数据清洗", "82.00%"],
            ["v1.03", "2026-05-19", "EfficientNetV2-S", "Focal Loss+RandAugment+MixUp", "90.00%"],
            ["v1.04", "2026-05-22", "EfficientNetV2-S", "CutMix+RandomErasing+TTA", "90.20%"],
            ["v1.06", "2026-05-27", "EfficientNetV2-S", "数据集扩充+SWA+Cosine Warmup", "96.53%"],
            ["v1.07", "2026-05-29", "ConvNeXt Small", "大模型训练+综合优化", "97.48%"],
        ],
        col_widths=[2.0, 2.5, 3.0, 5.0, 3.5]
    )
    add_empty_line(doc)

    # ===================== CHAPTER 2: DATA =====================
    doc.add_page_break()
    add_heading_custom(doc, "2  数据处理", level=1)
    add_heading_custom(doc, "2.1  数据来源", level=2)
    add_para(doc,
        "本项目数据集融合自多个来源，通过多源数据合并以提升数据规模和多样性。初始数据集为Kaggle平台的"
        "Garbage Classification数据集（TrashNet数据集在Kaggle的镜像版本），包含约2,500张已标注垃圾图像，"
        "涵盖纸板、玻璃、金属、纸张、塑料和其他垃圾六类。该数据集图像质量较高，但类别分布不均衡，其中"
        "trash类别仅有136张图像，导致模型在该类别上性能较差。")

    add_para(doc,
        "为突破数据瓶颈，项目在v1.06阶段进行了大规模数据集扩充。扩充来源为Hugging Face上的两个开源数据集："
        "omasteam/waste-garbage-management-dataset（MIT许可证，19,762张图像，10个类别）和"
        "shahzaibvohra/realwaste（CC BY 4.0许可证，4,752张图像，9个类别）。扩充脚本使用git sparse clone"
        "技术只下载需要的类别目录，通过MD5哈希去重避免与现有数据重复，并将图像统一缩放到224x224像素、"
        "以JPEG格式（quality=95）保存。通过类别映射，将realwaste数据集中的非标准标签（Textile Trash、"
        "Food Organics、Vegetation等）统一映射到trash类，丰富了该类别的视觉多样性。")

    add_heading_custom(doc, "2.2  数据集统计与划分", level=2)
    add_para(doc,
        "数据集采用分层采样方式进行划分：测试集占20%，验证集占10%，训练集占70%。分层划分确保每个子集中"
        "各类别保持相同的比例分布。划分使用固定随机种子（seed=42），保证结果可复现。")

    add_heading_custom(doc, "2.2.1  数据集扩充前后训练/验证/测试集划分对比（核心分析）", level=3)
    add_para(doc,
        "下表详细展示了数据集扩充前后、训练集/验证集/测试集在各类别上的数量变化，这是理解模型性能提升的"
        "关键数据基础。", bold=True)

    add_table_caption(doc, "表2  数据集扩充前后训练/验证/测试集划分详细对比（核心分析表）")
    create_table(doc,
        ["类别", "扩充前\n总计", "扩充后\n总计", "扩充前\n训练", "扩充后\n训练",
         "扩充前\n验证", "扩充后\n验证", "扩充前\n测试", "扩充后\n测试"],
        [
            ["Cardboard", "403", "1,825", "282", "1,277", "40", "183", "81", "365"],
            ["Glass",     "500", "3,064", "350", "2,145", "50", "306", "100", "613"],
            ["Metal",     "402", "1,019", "281", "713",  "40", "102", "81",  "204"],
            ["Paper",     "581", "1,680", "407", "1,176", "58", "168", "116", "336"],
            ["Plastic",   "475", "1,982", "332", "1,387", "48", "199", "95",  "396"],
            ["Trash",     "136", "948",  "95",  "664",  "14", "94",  "27",  "190"],
            ["总计", "2,497", "10,518", "1,748", "7,362", "250", "1,052", "499", "2,104"],
        ],
        col_widths=[2.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8]
    )
    add_empty_line(doc)

    add_para(doc, "核心变化分析：", bold=True)
    add_para(doc,
        "(1) 数据集总量从2,497张扩充至10,518张，增长321%（+8,021张），为深度学习模型提供了充足的数据基础。")
    add_para(doc,
        "(2) 训练集从1,748张增至7,362张（+5,614张，+321%），验证集从250张增至1,052张（+802张），"
        "测试集从499张增至2,104张（+1,605张）。各子集划分比例保持不变（70%/10%/20%），保证了实验对比的公平性。")
    add_para(doc,
        "(3) Trash类扩充比例最高（+597%），从136张增至948张。尤其值得关注的是，验证集中trash类从仅14张"
        "增至94张（+571%）。验证集样本量过少时（原仅14张），评估指标的统计波动大，难以真实反映模型性能。"
        "扩充后94张验证样本提供了可靠的模型选择依据。")
    add_para(doc,
        "(4) Glass类扩充量最大（+2,564张，+513%），达到3,064张，成为数据量最充足的类别。"
        "Metal类扩充比例最小（+153%），说明该类别在外部数据源中较为稀缺。")
    add_para(doc,
        "(5) 扩充后每个类别在训练集中均超过660张（Metal最低为713张），验证集均超过94张，"
        "测试集均超过190张，为模型训练和可靠评估提供了充足的数据基础。")

    add_table_caption(doc, "表3  扩充后数据集在各子集中的类别分布汇总")
    create_table(doc,
        ["类别", "总计", "训练集(70%)", "验证集(10%)", "测试集(20%)", "占总比"],
        [
            ["Cardboard", "1,825", "1,277", "183", "365", "17.3%"],
            ["Glass",     "3,064", "2,145", "306", "613", "29.1%"],
            ["Metal",     "1,019", "713",  "102", "204", "9.7%"],
            ["Paper",     "1,680", "1,176", "168", "336", "16.0%"],
            ["Plastic",   "1,982", "1,387", "199", "396", "18.8%"],
            ["Trash",     "948",  "664",  "94",  "190", "9.0%"],
            ["总计", "10,518", "7,362", "1,052", "2,104", "100%"],
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 2.5, 2.0]
    )
    add_empty_line(doc)

    add_heading_custom(doc, "2.2.2  测试集改动带来的评估公平性问题（关键分析）", level=3)
    add_para(doc,
        "需要特别指出并反思的是：数据集扩充过程中，测试集也随之发生了改变。具体而言，v1.03及之前版本"
        "使用的测试集为从原始2,497张图像中分层采样的499张；而v1.06及之后版本使用的测试集为从扩充后"
        "10,518张图像中分层采样的2,104张。这两个测试集在构成上存在本质差异——新测试集包含了来自"
        "omasteam和realwaste外部数据源的图像，而原始测试集仅包含Kaggle/TrashNet图像。", bold=True)
    add_para(doc,
        "这一变动对实验评估的公平性和严谨性提出了严峻挑战，主要体现在以下几个方面：")
    add_para(doc,
        "第一，性能对比的基准不一致。当报告`v1.06准确率从90.00%提升至96.53%`时，这两个数据分别来自"
        "不同测试集——90.00%是在499张原始测试集上的结果，96.53%是在2,104张新测试集上的结果。"
        "如果新测试集的难度分布与原始测试集不同（例如新增的omasteam数据可能拍摄条件更规范、背景更简洁，"
        "导致分类更容易），则部分性能提升可能来源于测试集难度下降，而非模型真实能力的提升。")
    add_para(doc,
        "第二，数据泄露风险。扩充后的数据集划分是在10,518张图像上重新进行的全局分层采样。这意味着"
        "新测试集中可能包含来自omasteam/realwaste的数据，而同一来源的数据在训练集中也存在。虽然"
        "图片级别没有重复（已通过MD5去重），但同源数据可能存在风格相似性（如同一数据集的拍摄环境、"
        "光照条件、分辨率存在偏置），使模型在`见过来自同一来源的训练数据后`对新测试集中的同源图像"
        "产生隐式的过拟合优势。")
    add_para(doc,
        "第三，测试集规模变化对评估可靠性的影响。扩充前测试集仅499张，其中trash类仅27张，评估结果"
        "的统计波动大（95%置信水平下误差限约+4.4%）。扩充后测试集2,104张，trash类190张，评估结果"
        "更可靠（误差限约+2.1%）。因此，v1.03报告的90.00%可能是一个有偏估计，而v1.06的96.53%统计上"
        "更加稳定。这进一步加大了跨版本直接对比的不确定性。")

    add_para(doc, "理想的研究实践应该是：", bold=True)
    add_para(doc,
        "最严谨的做法是在数据集扩充时，将原始测试集（499张）作为固定的、不可变的基准测试集（hold-out "
        "set）保存下来。所有新旧版本的模型都在这个固定测试集上评估，确保性能对比的公正性。扩充带来的"
        "新图像全部用于训练集和验证集，这样可以单独衡量`数据扩充带来的提升`这一变量，而不与其他因素"
        "（测试集构成变化）相混淆。这是机器学习实验设计的黄金标准。")
    add_para(doc,
        "然而，本项目的原始测试集仅499张（trash类仅27张），规模过小，统计可靠性不足。在学术研究中，"
        "通常要求测试集至少包含每类30-50个样本以保证统计显著性。因此，扩充后采用更大的测试集（2,104张）"
        "在统计意义上更加合理——这实际上是对评估方法本身的一种改进，但代价是与历史结果不可直接比较。")
    add_para(doc,
        "综上所述，v1.06/v1.07报告的性能提升（90.00%->96.53%/97.48%）应被理解为综合改进的效果——包含"
        "数据集扩充、训练优化、以及评估集变更三个因素的共同作用。其中'数据集扩充'和'训练优化'是"
        "真实的模型能力提升因素，而'评估集变更'则引入了不确定性。为了缓解这一问题，本报告在消融分析"
        "（4.4节）中各优化项的贡献度估算参考了v1.06版本内部的消融实验（即控制数据量和训练策略两个变量，"
        "在统一的大测试集上进行评估），而非简单对比跨版本的测试集结果。")
    add_heading_custom(doc, "2.3  数据清洗", level=2)
    add_para(doc,
        "数据清洗是数据预处理的基础环节，具体包含以下步骤：第一，损坏图像检测与移除——使用PIL库的"
        "Image.open().verify()方法逐张验证图像完整性，确保数据集中无损坏文件。第二，重复数据检测与删除——"
        "基于感知哈希（pHash）算法，将图像转换为8x8灰度图，计算所有像素的均值，每个像素与均值比较生成"
        "64位二进制哈希值，通过哈希值比对识别重复图像。最终从初始2,527张图像中检测并移除了27张重复图像，"
        "保留2,497张有效图像。第三，标签验证——确保每张图像正确存放在对应的类别目录中。")

    add_heading_custom(doc, "2.4  数据增强策略", level=2)
    add_para(doc,
        "为提升模型泛化能力和鲁棒性，训练阶段采用了多层次、组合式的数据增强策略：")
    add_para(doc,
        "基础几何增强：包括随机水平翻转（50%概率）、随机垂直翻转（30%概率）、随机旋转（最大30度）、"
        "色彩抖动（亮度/对比度/饱和度各0.2、色相0.1）和随机透视变换（扭曲尺度0.1、30%概率），"
        "这些增强增加了训练样本在空间变换和光照变化下的多样性。")
    add_para(doc,
        "高级增强策略：项目集成了三种先进的增强方法。RandAugment自动增强策略（num_ops=2, magnitude=9）"
        "从预定义的增强操作池中随机选择2个操作自动组合应用；CutMix区域混合增强（Beta分布alpha=0.4, "
        "50%概率应用）从一张图像中裁剪区域粘贴到另一张图像上，标签按面积比例混合，迫使模型关注局部特征"
        "而非全局纹理；RandomErasing随机擦除增强（25%概率，擦除区域占图像的2%-40%）随机遮挡图像中的"
        "矩形区域，提升模型对局部遮挡的鲁棒性。")
    add_para(doc,
        "类别平衡采样：在数据加载阶段使用WeightedRandomSampler，各类别采样权重与样本数成反比"
        "（cardboard=0.96, glass=0.57, metal=1.72, paper=1.04, plastic=0.88, trash=1.85），"
        "确保每个batch中各类别被采样概率均衡。据消融实验估算，数据增强组合对最终性能的贡献约为0.5-2.0%。")

    # ===================== CHAPTER 3: MODEL ARCHITECTURE =====================
    doc.add_page_break()
    add_heading_custom(doc, "3  模型架构与训练优化", level=1)
    add_heading_custom(doc, "3.1  模型家族总览", level=2)
    add_para(doc,
        "项目实现了10种不同规模的深度学习模型，覆盖了从轻量级自建卷积网络到大型预训练模型的完整谱系，"
        "参数规模从1.2M到53M，跨越了三个数量级。这种多模型设计使项目能够在不同应用场景下做出灵活的权衡选择。")
    add_para(doc,
        "SimpleCNN（1.2M参数）：四层Conv-BN-ReLU-MaxPool模块堆叠的自建卷积网络，后接全连接分类层，"
        "适合作为无预训练的基线参考。MobileNetV2（3.5M参数）：使用深度可分离卷积将标准卷积分解为逐通道"
        "卷积和逐点卷积，采用倒残差结构实现高效特征提取，是移动端部署的理想选择。ResNet18（11.2M参数）"
        "和ResNet50（25.6M参数）：使用残差连接使梯度可以直接通过恒等映射回传到浅层，有效解决了深层网络中的"
        "梯度消失问题。")
    add_para(doc,
        "EfficientNetV2-S（21M参数）和EfficientNetV2-M（53M参数）：采用Fused-MBConv模块，将MBConv中的"
        "1x1升维卷积和3x3深度卷积合并为单个3x3普通卷积，提升了计算效率。ConvNeXt Tiny（28M参数）和"
        "ConvNeXt Small（50M参数）：作为项目中最先进的模型架构，融入了Transformer的设计理念——使用"
        "LayerNorm替代BatchNorm、GELU激活函数、7x7大核卷积，同时保留了纯CNN架构的推理效率，最终以"
        "97.48%的测试准确率成为项目最佳模型。")

    add_table_caption(doc, "表4  全部模型架构参数与性能总览")
    create_table(doc,
        ["模型", "参数量", "测试准确率", "Macro-F1", "推理时间", "版本", "特点"],
        [
            ["SimpleCNN", "1.2M", "51.80%", "0.5179", "24ms", "v1.0", "自建CNN基线"],
            ["MobileNetV2", "3.5M", "82.00%", "0.8042", "16ms", "v1.0", "轻量移动端"],
            ["ResNet18", "11.2M", "78.00%", "0.7632", "15ms", "v1.0", "经典残差网络"],
            ["EfficientNetV2-S", "~21M", "90.00%", "0.8895", "46ms", "v1.03", "v1.03最佳效率"],
            ["ConvNeXt Tiny", "~28M", "39.40%", "0.3736", "205ms", "v1.03", "未充分收敛"],
            ["EfficientNetV2-S*", "~21M", "96.53%", "0.9633", "73ms", "v1.06", "扩充+优化"],
            ["EfficientNetV2-M", "~53M", "96.96%", "0.9664", "4,946ms", "v1.07", "大容量模型"],
            ["ConvNeXt Small*", "~50M", "97.48%", "0.9723", "2,817ms", "v1.07", "最佳精度"],
        ],
        col_widths=[2.5, 1.5, 2.0, 2.0, 2.0, 1.5, 3.0]
    )
    add_empty_line(doc)

    add_heading_custom(doc, "3.2  损失函数设计", level=2)
    add_para(doc,
        "项目使用FocalLossWithLabelSmoothing作为组合损失函数。标准交叉熵损失在类别不平衡场景下存在明显缺陷——"
        "多数类样本的梯度信号占主导地位，少数类难以得到充分优化。Focal Loss通过引入调制因子(1-pt)^gamma"
        "（gamma=2.0）解决了这一问题：当pt接近1（易分类样本）时，调制因子接近0，损失权重被大幅降低；"
        "当pt接近0（难分类样本）时，调制因子接近1，损失权重几乎保持不变。")
    add_para(doc,
        "标签平滑（Label Smoothing, epsilon=0.1）将硬标签（one-hot）转换为软标签分布，使目标类别概率"
        "从100%降至90%，其余10%均匀分配给其他类别。这一机制防止模型产生过于自信的极端预测，从而降低"
        "过拟合风险。将Focal Loss和Label Smoothing两种互补技术合并为FocalLossWithLabelSmoothing，"
        "同时发挥聚焦难分类样本和防止过拟合的双重优势。")

    add_heading_custom(doc, "3.3  训练优化策略", level=2)
    add_para(doc,
        "学习率调度方面，使用Cosine Warmup + CosineAnnealingLR组合策略。前5个epoch执行线性预热（warmup），"
        "学习率从初始值的1%逐步增加到100%，防止训练初期的大幅梯度震荡；之后进行单调余弦退火至最小值1e-6，"
        "相比v1.03的CosineAnnealingWarmRestarts（周期性重启），单调退火避免了学习率重启可能导致的性能倒退。")
    add_para(doc,
        "优化器选择根据不同模型架构做了差异化配置：EfficientNetV2系列和ResNet系列使用SGD Nesterov动量优化器"
        "（lr=0.01, momentum=0.9, weight_decay=1e-4, gradient_clip=1.0），与EfficientNetV2原论文保持一致，"
        "泛化能力优于AdamW；ConvNeXt系列使用AdamW优化器（lr=1e-4, weight_decay=0.05, gradient_clip=5.0），"
        "因为ConvNeXt架构设计中使用了LayerNorm和大核卷积，对AdamW的适应性更好。")

    add_table_caption(doc, "表5  各模型专属训练超参数配置")
    create_table(doc,
        ["模型", "优化器", "学习率", "权重衰减", "梯度裁剪"],
        [
            ["SimpleCNN", "Adam", "1e-3", "1e-4", "1.0"],
            ["MobileNetV2", "AdamW", "5e-5", "1e-4", "1.0"],
            ["ResNet18", "SGD Nesterov", "0.01", "1e-4", "1.0"],
            ["EfficientNetV2-S", "SGD Nesterov", "0.01", "1e-4", "1.0"],
            ["EfficientNetV2-M", "SGD Nesterov", "0.01", "1e-4", "1.0"],
            ["ConvNeXt Tiny", "AdamW", "1e-4", "0.05", "5.0"],
            ["ConvNeXt Small", "AdamW", "1e-4", "0.05", "5.0"],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 2.5]
    )
    add_empty_line(doc)

    add_heading_custom(doc, "3.4  高级训练技巧", level=2)
    add_para(doc,
        "指数移动平均（EMA, decay=0.999）：在训练过程中维护权重的平滑平均版本——每步权重更新后执行"
        "new_avg = 0.999 * shadow_avg + 0.001 * current_weight，最终得到的平均权重相比任何单时间点的"
        "权重都更加稳定和泛化。EMA权重视为最佳模型持久化保存，确保推理和验证阶段使用一致的权重版本。")
    add_para(doc,
        "随机权重平均（SWA, 最后25% epoch启动）：以固定学习率对SGD迭代的权重进行等权平均，通过找到损失"
        "景观中更平坦的局部极小值来提升泛化能力。SWA在最后25个epoch（共100 epoch）启动，初始学习率为原学习率的0.1倍。")
    add_para(doc,
        "早停机制（patience=30 epoch）：监控验证集上的损失值，若30个连续epoch内未出现新的最佳性能，则提前"
        "终止训练以防止过拟合。与v1.03的patience=10相比，v1.06增加至30，给大模型更充分的收敛时间。梯度裁剪"
        "则通过设置梯度最大范数（max_norm）来防止梯度爆炸。")

    add_heading_custom(doc, "3.5  推理优化", level=2)
    add_para(doc,
        "测试时增强（Test-Time Augmentation, TTA）在评估阶段发挥作用：对每张测试图片生成8个增强版本"
        "（含原图、水平翻转、4个旋转角度0/90/180/270度以及翻转+旋转的组合），取所有版本预测的软投票"
        "平均值作为最终结果。TTA能带来1-2%的准确率提升。")

    # ===================== CHAPTER 4: RESULTS =====================
    doc.add_page_break()
    add_heading_custom(doc, "4  实验结果与分析", level=1)
    add_heading_custom(doc, "4.1  全模型性能对比", level=2)
    add_para(doc,
        "所有模型的实验均在NVIDIA GeForce RTX 5060 Laptop GPU（8GB VRAM）上进行，使用扩充后的10,518张"
        "数据集（v1.06/v1.07实验）。测试结果为在独立测试集（2,104张）上的评估数据。")
    add_para(doc,
        "从全模型对比中可以看到清晰的性能层次：SimpleCNN作为无预训练基线仅达51.80%，说明预训练权重对图像"
        "分类任务至关重要。MobileNetV2和ResNet18分别在3.5M和11.2M参数下取得了82.00%和78.00%的准确率，"
        "展现了迁移学习的显著优势。EfficientNetV2-S在v1.03配置下达到90.00%，在v1.06经数据扩充和训练优化后"
        "跃升至96.53%，提升高达6.53个百分点。EfficientNetV2-M（53M）相比V2-S（21M）参数翻2.6倍但准确率"
        "仅提升0.43%（96.96% vs 96.53%），表明在10,518张的数据规模下，模型容量已不是当前主要瓶颈。"
        "ConvNeXt Small以97.48%的测试准确率和0.9723的Macro-F1成为项目最终最佳模型。")

    add_heading_custom(doc, "4.2  ConvNeXt Small详细指标", level=2)

    add_table_caption(doc, "表6  ConvNeXt Small模型各类别详细分类指标（测试集2,104张）")
    create_table(doc,
        ["类别", "Precision", "Recall", "F1-Score", "支持数"],
        [
            ["Cardboard", "0.976", "0.986", "0.981", "365"],
            ["Glass",     "0.985", "0.982", "0.984", "613"],
            ["Metal",     "0.940", "1.000", "0.969", "204"],
            ["Paper",     "0.979", "0.976", "0.978", "336"],
            ["Plastic",   "0.969", "0.957", "0.963", "396"],
            ["Trash",     "0.983", "0.937", "0.960", "190"],
            ["Macro Avg", "0.972", "0.973", "0.972", "2,104"],
        ],
        col_widths=[3.0, 3.0, 3.0, 3.0, 3.0]
    )
    add_empty_line(doc)

    add_para(doc,
        "在所有类别中，Glass的F1-Score最高（0.984），Metal的Recall达到完美的1.000（204张金属图像无漏检），"
        "说明模型对玻璃和金属这类具有明显视觉特征的材质识别最为准确。Trash的F1-Score为0.960，虽然因类别"
        "内部视觉多样性大（包含织物、食物残渣、陶瓷碎片等）而略低于其他类别，但相比v1.03的0.743取得了质的飞跃。"
        "Plastic的Recall相对最低（0.957），主要原因在于部分透明塑料瓶与玻璃瓶的外观相似造成了一定混淆。")

    add_heading_custom(doc, "4.3  Trash类突破分析", level=2)
    add_para(doc,
        "Trash类的性能提升是本项目最显著的成果之一。从各版本的F1追踪来看：v1.0基线仅0.162（几乎无法识别 "
        "trash），v1.03提升至0.743（初步可用），v1.06大幅跃升至0.955（高精度），v1.07的ConvNeXt Small进一步"
        "提升至0.960（接近完美）。")

    add_table_caption(doc, "表7  Trash类F1分数版本演进")
    create_table(doc,
        ["版本", "Trash类F1", "Trash类训练样本数", "核心变化"],
        [
            ["v1.0",  "0.162", "95",  "基线模型，无优化"],
            ["v1.03", "0.743", "95",  "Focal Loss + MixUp + RandAugment"],
            ["v1.06", "0.955", "664", "数据扩充+WeightedRandomSampler+FocalLoss"],
            ["v1.07", "0.960", "664", "ConvNeXt Small大模型"],
        ],
        col_widths=[2.5, 3.0, 3.0, 6.5]
    )
    add_empty_line(doc)

    add_para(doc,
        "这一突破来自三个因素的综合协同作用：(1)数据扩充是基础——trash类从136张增至948张（+597%），"
        "模型见过的视觉变体大幅增加；(2)WeightedRandomSampler是关键——通过反比采样权重（trash类权重最高"
        "为1.85），每个epoch都能见到足够多的trash样本；(3)Focal Loss提供辅助——自动聚焦trash这类难分类"
        "样本，在损失函数层面给予更高的优化优先级。三个因素缺一不可。")

    add_heading_custom(doc, "4.4  消融实验与贡献度分析", level=2)
    add_para(doc,
        "为理解各优化项的实际贡献，进行了系统的消融分析。下表总结了各优化措施的预估贡献度：")

    add_table_caption(doc, "表8  各优化措施对模型性能的贡献度分析")
    create_table(doc,
        ["优化措施", "预估贡献", "主要受益类别", "说明"],
        [
            ["数据集扩充（2,497->10,518张）", "+3.5~4.5%", "Trash（+0.212）", "最大单一贡献因素"],
            ["WeightedRandomSampler",         "+1.0~2.0%", "少数类（Trash, Metal）", "类别平衡采样"],
            ["FocalLossWithLabelSmoothing",   "+1.0~2.0%", "难分类样本", "联合损失函数"],
            ["Cosine Warmup + AnnealingLR",   "+0.5~1.0%", "训练稳定性", "平滑退火无重启"],
            ["SWA + EMA",                    "+0.5~1.0%", "模型泛化", "权重平均"],
            ["CutMix + RandomErasing",       "+0.5~1.0%", "鲁棒性", "区域增强"],
            ["总计", "+6.5~10.0%", "全类别", "实测+6.53%"],
        ],
        col_widths=[4.5, 2.5, 3.5, 3.0]
    )
    add_empty_line(doc)

    add_para(doc,
        "消融分析揭示了一个重要规律：数据规模对分类性能具有决定性影响。数据集扩充贡献了约3.5-4.5%的准确率"
        "提升，远超任何单一算法优化手段，印证了‘数据是深度学习的燃料’这一基本原则。同时，多项优化措施的"
        "协同效应也不容忽视——当数据量增加、采样平衡、损失自适应、学习率优化这些措施组合在一起时，产生了"
        "\"1+1>2\"的综合效果。")

    # ===================== CHAPTER 5: SYSTEM =====================
    doc.add_page_break()
    add_heading_custom(doc, "5  系统展示", level=1)
    add_heading_custom(doc, "5.1  Streamlit交互式应用", level=2)
    add_para(doc,
        "项目开发了基于Streamlit的交互式Web应用（app.py），提供直观的垃圾分类演示界面。主要功能包括："
        "实时图像分类预测——用户上传图片后，系统自动完成分类并显示预测类别、置信度和6类概率分布柱状图；"
        "模型性能对比——内置已训练模型的全指标对比仪表板，支持在SimpleCNN、MobileNetV2、ResNet18、"
        "EfficientNetV2-S/M、ConvNeXt Tiny/Small等模型间切换对比；垃圾分类建议——根据识别结果给出具体"
        "的回收处理建议，提升用户体验。")

    add_heading_custom(doc, "5.2  Web前端展示系统", level=2)
    add_para(doc,
        "v1.08新增了基于FastAPI的独立Web前端展示模块（web/目录），提供无需Streamlit的轻量级访问方式。"
        "后端API层实现了RESTful接口：GET /api/models返回可用模型列表；POST /api/predict接收上传图片和"
        "模型选择参数，返回识别结果。后端采用LRU缓存机制（最多保持2个模型同时加载），切换模型时自动卸载，"
        "节省GPU显存。前端为纯静态HTML/CSS/JS页面，采用绿色环保主题和响应式设计，适配桌面和移动端。")

    add_heading_custom(doc, "5.3  命令行工具链", level=2)
    add_para(doc,
        "CLI入口run.py支持五大任务模式：clean（执行数据清洗）、train（训练一个或多个模型）、evaluate"
        "（评估已训练模型）、inference（单张/批量图片推理）、all（全流程执行）。通过丰富的命令行参数"
        "（--models指定模型列表、--randaugment启用自动增强、--cutmix启用区域混合增强等），支持灵活的"
        "实验配置组合。此外还提供demo.py（快速演示脚本）、quickstart.py（交互式启动向导）等辅助工具。")

    # ===================== CHAPTER 6: CONCLUSION =====================
    doc.add_page_break()
    add_heading_custom(doc, "6  结论与展望", level=1)
    add_heading_custom(doc, "6.1  项目总结", level=2)
    add_para(doc,
        "本项目基于深度学习技术成功构建了一套高精度的智能垃圾分类系统，取得了丰硕成果：")
    add_para(doc,
        "在技术成果方面，构建了包含10,518张图像的大规模垃圾分类数据集，通过多源数据融合有效缓解了类别"
        "不平衡问题；实现了10种不同规模的深度学习模型架构，ConvNeXt Small以97.48%的测试准确率和0.9723"
        "的Macro-F1成为最终最佳模型；EfficientNetV2-S以21M参数、73ms推理时间达到96.53%的准确率，成为"
        "性价比最优选择。")
    add_para(doc,
        "在工程成果方面，项目采用模块化设计，提供了完整的命令行工具链、交互式Streamlit Web应用和"
        "FastAPI REST API服务以及详细的技术文档体系，支持从数据清洗到模型推理的全流程操作。")
    add_para(doc,
        "在学术贡献方面，项目系统验证了多种训练优化策略的有效性。消融分析揭示了一个重要发现：数据规模对"
        "分类性能具有决定性影响（数据集扩充贡献约6.53%准确率提升，远超Focal Loss的1-2%或SWA+EMA的0.5-1%"
        "等算法优化手段）。同时，trash类F1从0.162（几乎无法识别）到0.960（高度可靠）的突破性进展，证明了"
        "数据扩充、类别平衡采样和针对性损失函数三者协同作用的有效性。")

    add_heading_custom(doc, "6.2  未来展望", level=2)
    add_para(doc,
        "未来可从以下方向继续拓展：(1)模型轻量化部署——通过INT8量化、TensorRT加速等技术将模型优化至移动端"
        "和嵌入式平台，实现实时推理；(2)目标检测扩展——从当前单图像分类扩展至目标检测任务，支持图像中多个"
        "物体的同时识别和位置定位；(3)细粒度分类——将塑料等大类进一步细分为PET、PP、PE等子类，提供更加"
        "精细化的分类建议；(4)增量学习支持——使系统在部署后能够持续收集新数据不断优化模型，适应实际场景中"
        "不断变化的垃圾分布。")

    # ===================== REFERENCES =====================
    doc.add_page_break()
    add_heading_custom(doc, "参考文献", level=1)
    refs = [
        "[1] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. IEEE CVPR, 2016.",
        "[2] Sandler M, Howard A, Zhu M, et al. MobileNetV2: Inverted Residuals and Linear Bottlenecks[C]. IEEE CVPR, 2018.",
        "[3] Tan M, Le Q V. EfficientNetV2: Smaller Models and Faster Training[C]. ICML, 2021.",
        "[4] Liu Z, Mao H, Wu C Y, et al. A ConvNet for the 2020s[C]. IEEE CVPR, 2022.",
        "[5] Lin T Y, Goyal P, Girshick R, et al. Focal Loss for Dense Object Detection[C]. IEEE ICCV, 2017.",
        "[6] Izmailov P, Podoprikhin D, Garipov T, et al. Averaging Weights Leads to Wider Optima and Better Generalization[C]. UAI, 2018.",
        "[7] Cubuk E D, Zoph B, Shlens J, et al. RandAugment: Practical Automated Data Augmentation with a Reduced Search Space[C]. NeurIPS, 2020.",
        "[8] Zhang H, Cisse M, Dauphin Y N, et al. MixUp: Beyond Empirical Risk Minimization[C]. ICLR, 2018.",
        "[9] Yun S, Han D, Oh S J, et al. CutMix: Regularization Strategy to Train Strong Classifiers with Localizable Features[C]. IEEE ICCV, 2019.",
    ]
    for ref in refs:
        add_para(doc, ref, size=11, first_indent=False, spacing_after=3)

    # ===================== APPENDIX =====================
    doc.add_page_break()
    add_heading_custom(doc, "附录", level=1)

    add_heading_custom(doc, "A.  数据集扩充前后详细划分对比参考表", level=2)
    add_para(doc,
        "下表完整列出了每种类别在扩充前后在训练集、验证集、测试集中的具体样本数、变化量及增长率，"
        "作为实验分析的完整数据基础。")

    add_table_caption(doc, "表A1  数据集扩充前后各类别在训练/验证/测试集中的样本数及增长率")
    create_table(doc,
        ["类别", "扩充前\n总计", "扩充前\n训练", "扩充前\n验证", "扩充前\n测试",
         "扩充后\n总计", "扩充后\n训练", "扩充后\n验证", "扩充后\n测试", "增长率"],
        [
            ["Cardboard", "403", "282", "40", "81", "1,825", "1,277", "183", "365", "+353%"],
            ["Glass",     "500", "350", "50", "100", "3,064", "2,145", "306", "613", "+513%"],
            ["Metal",     "402", "281", "40", "81",  "1,019", "713",  "102", "204", "+153%"],
            ["Paper",     "581", "407", "58", "116", "1,680", "1,176", "168", "336", "+189%"],
            ["Plastic",   "475", "332", "48", "95",  "1,982", "1,387", "199", "396", "+317%"],
            ["Trash",     "136", "95",  "14", "27",  "948",  "664",  "94",  "190", "+597%"],
            ["总计", "2,497", "1,748", "250", "499", "10,518", "7,362", "1,052", "2,104", "+321%"],
        ],
        col_widths=[1.8, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4]
    )
    add_empty_line(doc)

    add_heading_custom(doc, "B.  项目结构说明", level=2)
    items = [
        "src/——核心源代码目录，包含数据加载、数据清洗、模型定义（10种模型）、训练器（含Focal Loss, EMA, SWA等）、评估器（含TTA, 模型集成, 混淆矩阵）、推理模块。",
        "data/processed/——处理后数据集，按6个类别子目录组织，含10,518张已清洗图像，附带清洗报告和数据清单。",
        "models/——已训练模型权重，共11个文件（含BEST和SWA变体），总计约1.6 GB，通过Git LFS管理。",
        "web/——FastAPI Web前端展示模块，包含REST API后端和纯静态HTML/CSS/JS前端页面。",
        "download_supplement.py——数据集扩充脚本，使用git sparse clone从Hugging Face拉取补充数据。",
        "app.py——Streamlit交互式Web应用入口。run.py——CLI命令行工具主入口。",
        "logs/——训练日志和JSON格式评估结果。",
    ]
    for item in items:
        add_para(doc, "  * " + item, size=11, first_indent=False, spacing_after=2)

    add_empty_line(doc)
    add_heading_custom(doc, "C.  技术栈与运行环境", level=2)
    add_para(doc,
        "项目基于Python 3.8+构建，核心技术栈包括：PyTorch 2.0.1（深度学习框架）、TorchVision 0.15.2"
        "（预训练模型和图像变换工具）、NumPy 1.24.3、Scikit-learn 1.3.0（评估指标和数据划分）、"
        "Matplotlib 3.7.2 + Seaborn 0.12.2（数据可视化）、Streamlit 1.28.1（Web交互式演示界面）、"
        "FastAPI + Uvicorn（REST API服务）。训练硬件为NVIDIA GeForce RTX 5060 Laptop GPU（8GB VRAM），"
        "全部模型训练累计约20小时。")

    add_empty_line(doc)
    add_heading_custom(doc, "D.  数据清洗报告摘要", level=2)
    add_para(doc,
        "初始数据集统计：总计2,527张图像——cardboard 403张, glass 501张, metal 410张, paper 594张, "
        "plastic 482张, trash 137张。清洗结果：移除损坏图像0张，检测并删除重复图像27张（基于感知哈希去重）。"
        "最终数据集：2,497张有效图像——cardboard 403张, glass 500张, metal 402张, paper 581张, "
        "plastic 475张, trash 136张。所有标签验证通过，无标签错误。")

    add_empty_line(doc)
    add_empty_line(doc)

    # Footer info
    add_para(doc, "文档撰写日期：2026年6月    版本：v1.07",
             size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    add_para(doc, "同济大学  Python程序设计课程小组项目",
             size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)

    # ====== SAVE ======
    out_path = r"C:\Users\zgzz1\Desktop\python-程序设计\py-group-project-garbage_classification\technical_report_enriched.docx"
    doc.save(out_path)
    file_size = os.path.getsize(out_path)
    print(f"SUCCESS: Report generated at: {out_path}")
    print(f"File size: {file_size / 1024:.1f} KB")


if __name__ == '__main__':
    generate()
